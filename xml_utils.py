from lxml import etree
from docx import Document
from docx.shared import Pt
from io import BytesIO
from collections import OrderedDict
import re
from wtforms import FieldList, FormField # Added imports

# Merged FORM_LABEL_MAP with balises from saisine-2.20-unformatted.txt
# Prioritizing entries from the Perl script for numbering and titles
FORM_LABEL_MAP = {
    "Formulaire_Apafis": "Formulaire APAFIS",
    "InformationsGenerales": "1. Informations Générales",
    "TitreProjet": "1.2. Titre du projet :",
    "NumVersion": "Numéro de version :",
    "ReferenceDossier": "1.1. Référence Dossier :",
    "DureeProjet": "1.3. Durée du projet :",
    "DureeAnnees": "Années :",
    "DureeMois": "Mois :",
    "DureeJours": "Jours :",
    "DebutQuandProjetAutorise": "Dès que le projet est autorisé :",
    "DatePrevueDebutProjet": "1.4. Date prévue de début du projet (Année-Mois-Jour):",
    "Date": "Date :",
    "RNT": "2. Rnt (inutilisé)",
    "InformationsAdministrativesEtReglementaires": "3. Informations Administratives et Réglementaires",
    "EtablissementUtilisateur": "3.1. L établissement utilisateur (EU)",
    "AgrementUE": "3.1.1. Agrément de l'EU où seront utilisés les animaux :",
    "NomUE": "Nom :",
    "NumeroAgrement": "Numéro d'agrément :",
    "DateDelivrance": "Date de délivrance de l'agrément :",
    "CiviliteResponsable": "Civilité :",
    "NomResponsable": "Nom du responsable :",
    "PrenomResponsable": "Prénom du responsable :",
    "EMailResponsable": "Adresse électronique du responsable :",
    "CiviliteDelegataire": "Civilité du délégataire :",
    "NomDelegataire": "Nom de la personne délégataire du responsable présente dans l'EU :",
    "PrenomDelegataire": "Prénom de la personne délégataire du responsable présente dans l'EU :",
    "EMailDelegataire": "Adresse électronique de ce délégataire :",
    "ComiteEthiqueAgree": "3.1.2. Comité d'éthique agréé par le MESR dont relève l'EU :",
    "ResponsablesMiseEnOeuvre": "3.1.3. Responsable(s) de la mise en oeuvre générale du projet dans l'EU et de sa conformité à l'autorisation de projet :",
    "CoordonneesResponsablesMiseEnOeuvre": "Responsable(s) :",
    "Civilite": "Civilité :",
    "Nom": "Nom :",
    "Prenom": "Prenom :",
    "AdressePostale": "Adresse postale :",
    "NomLaboratoire": "Nom du Laboratoire :",
    "ComplementAdresse": "Complément :",
    "NumeroRue": "Numéro de rue :",
    "Voie": "Voie :",
    "CodePostal": "Code Postal :",
    "Ville": "Ville :",
    "Pays": "Pays :",
    "Email": "Email :",
    "NumTelephone": "Téléphone :",
    "ResponsablesBienEtre": "3.1.4. Responsable(s) du bien-être des animaux :",
    "CoordonneesResponsablesBienEtre": "Responsable(s) :",
    "Personnel": "3.2. Le personnel ",
    "ConceptionProceduresExp": "Compétences des personnes participant au projet : - la conception des procédures expérimentales et des projets : ",
    "ApplicationProceduresExp": " - l'application de procédures expérimentales aux animaux : ",
    "SoinAuxAnimaux": " - les soins aux animaux : ",
    "MiseAMort": " - la mise à mort des animaux : ",
    "Projet": "3.3. Le projet",
    "ObjectifDuProjet": "3.3.1. L'objectif du projet :",
    "JustificationProjet": "Est-il :",
    "JustifieEducativement": " - justifié du point de vue éducatif ? : ",
    "RequisLoi": " - requis par la loi ? : ",
    "JustifieScientifiquement": " - justifié du point de vue scientifique ? : ",
    "InformationsJustification": "Quelle est l’instance qui a évalué l’intérêt de ce projet ? ",
    "DescriptionProjet": "3.3.2. Description du projet :",
    "DescriptionProjet2": "3.3.2. Description du projet :",
    "ObjectifsDuProjet": "3.3.2.1. Objectifs du projet :",
    "DerouleDuProjet": "3.3.2.2. Déroulé du projet :",
    "BeneficesDuProjet": "3.3.2.3. Bénéfices attendus du projet (version scientifique):",
    "NuisancesAnimaux": "3.3.2.4. Nuisances ou effets indésirables attendus sur les animaux (version scientifique):",
    "MethodeMiseAMort": "3.3.3. Précisez, le cas échéant, la ou les méthodes de mise à mort prévue(s) :",
    "ElementsScientifiquesJustifiantDemandeMiseAMort": "3.3.4. Précisez, le cas échéant, les éléments scientifiques justifiant la demande de dérogation concernant la méthode de mise à mort envisagée :",
    "StrategieExperimentations": "3.3.5. Stratégie d'expérimentation ou d'observation et approche statistique utilisée afin de réduire au minimum le nombre d'animaux, la douleur, la souffrance et l'angoisse, infligées et l'impact environnemental, le cas échéant . si une étude statistique est prévue, indiquez et justifiez les tests choisis (INUTILISE):",
    "Strategies3R": "3.3.6. Stratégies de Remplacement, de Réduction et de Raffinement",
    "Remplacement": "3.3.6.1. Remplacement",
    "Reduction": "3.3.6.2. Réduction",
    "Raffinement": "3.3.6.3. Raffinement",
    "Animaux": "3.4. Les animaux",
    "JustificationRecoursAuxAnimaux": "3.4.1. Justifiez la nécessité d'avoir recours à des animaux pour atteindre les objectifs du projet (inutilisé):",
    "ListeAnimaux": "3.4.2. Espèces animales ou types d'animaux utilisés (le champ 3.4.1. est supprimé) :",
    "Souris": "Souris (Mus musculus) [A1]",
    "Rats": "Rats (Rattus norvegicus) [A2]",
    "Cobayes": "Cobayes (Cavia porcellus) [A31]",
    "HamstersSyriens": "Hamsters (syriens) (Mesocricetus auratus) [A4]",
    "HamstersChinois": "Hamsters (chinois) (Cricetulus griseus) [A5]",
    "Gerbille": "Gerbilles de Mongolie (Meriones unguiculatus) [A6]",
    "AutresRongeurs": "Autres rongeurs (Rodentia) [A7]",
    "Lapins": "Lapins (Oryctolagus cuniculus) [A8]",
    "Chats": "Chats (Felis catus) [A9]",
    "Chiens": "Chiens (Canis familiaris) [A10];",
    "Furets": "Furets (Mustela putorius furo) [A11]",
    "AutresCarnivores": "Autres carnivores (carnivora) [A12]",
    "Chevaux": "Chevaux, ânes et croisements (Equidae) [A13]",
    "Porcs": "Porcs (Sus scrofa domesticus) [A14]",
    "Caprins": "Caprins (Capra aegagrus hircus) [A15]",
    "Ovins": "Ovins (Ovis aries) [A16]",
    "Bovins": "Bovins (Bos taurus) [A17]",
    "Prosimien": "Prosimiens (prosimia) [A18]",
    "Ouistitis": "Ouistitis et tamarins (par exemple, Callithrix jacchus) [A19]",
    "SingeCynomologue": "Singe cynomolgus (Macaca fascicularis) [A20]",
    "SingeRhesus": "Singe rhésus (Macaca mulatta) [A21]",
    "Vervets": "Vervets (Chlorocebus spp.) (soit pygerythrus, soit sabaeus) [A22]",
    "Babouins": "Babouins (Papio son.) [A23]",
    "Sairimis": "Saïmiris (par exemple, Saimiri sciureus) [A24]",
    "AutresSingesAncienMonde": "Autres espèces de singes de l'Ancien Monde (autres espèces de Cercopithecoidea)[A25-1]",
    "AutresSingesNouveauMonde": "Autres espèces de singes du Nouveau Monde[A25-2]",
    "AutresPrimateNonHumain": "Autres espèces de primates non humains [cocher [A25-1] ou [A25-2]]",
    "SingesAntropoides": "SingesAntropoides",
    "AutresMammiferes": "Autres mammifères (autres Mammalia) [A27]",
    "PoulesDomestiques": "Poules domestiques (Gallus gallus domesticus) [A28]",
    "Dindes": "Dindes",
    "AutresOiseaux": "Autres oiseaux (autres Aves) [A29]",
    "Reptiles": "Reptiles (Reptilia) [A30]",
    "GrenouillesRana": "Grenouilles Rana (Rana temporaria et Rana pipiens) [ABI]",
    "GrenouillesXenopus": "Grenouilles Xenopus (Xenopus laevis et Xenopus tropicalis) [A32]",
    "AutresAmphibiens": "Autres amphibiens (autres Amphibia) [A33]",
    "PoissonsZebres": "Poissons zèbres (Danio rerio) [A34]",
    "Bars": "Bars",
    "Saumons": "Saumons",
    "Guppies": "Guppies",
    "AutresPoissons": "Autres poissons (autres Pisces) [A35]",
    "Cephalopodes": "Céphalopodes (Cephalopoda) [A36]",
    "PertinenceAnimauxChoisis": "3.4.3. Justifiez la pertinence de l'(des) espèce(s) animale(s) choisie(s) :",
    "EspecesMenacees": "3.4.4. S'agit-il de spécimens d'espèces menacées énumérées à l'annexe A du règlement (CE) n° 338/97 du Conseil du 9 décembre 1996 relatif à la protection des espèces de faune et de flore sauvages par le contrôle et leur commerce ?",
    "OuiEspecesMenacees": "Si oui, éléments scientifiques démontrant que la finalité de la procédure expérimentale ne peut être atteinte en utilisant d'autres espèces que celles énumérées dans cette annexe :",
    "Primate": "3.4.5. S'agit-il de spécimens de primates non humains ?",
    "OuiPrimate": "Si oui, éléments scientifiques démontrant que la finalité de la procédure expérimentale ne peut être atteinte en utilisant d'autres espèces de primates non humains:",
    "CaptureNature": "3.4.6. S'agit-il d'animaux capturés dans la nature ? ",
    "OuiCaptureNature": "Si oui, éléments scientifiques démontrant que la finalité de la procédure expérimentale ne peut être atteinte en utilisant d'autres animaux que ceux capturés dans la nature :",
    "EspeceDomestique": "3.4.7. S'agit-il d'animaux d'espèces domestiques, errants ou vivant à l'état sauvage ?",
    "OuiEspeceDomestique": "Si oui, éléments scientifiques démontrant que la finalité de la procédure expérimentale ne peut être atteinte qu'en utilisant ces animaux.",
    "CategorieAnimauxUtilises": "3.4.8. Catégorie des animaux utilisés dans le projet :",
    "AnimauxEnCaptivite": " - Animaux tenus en captivité (domestiques ou non domestiques) : ",
    "AnimauxNonDomNonCaptifs": " - Animaux non domestiques non tenus en captivité : ",
    "ReferencesDerogationsEspProtegee": "Si les animaux utilisés sont des spécimens d'espèces protégées en application de l'article L. 411-1 du Code de l'environnement, indiquez les références de la dérogation accordée pour effectuer la capture des animaux dans le milieu naturel (4° de l'article L. 411-2 du Code de l'environnement) :",
    "ReferencesDerogationsChasseAutorisee": "Si les animaux utilisés sont des spécimens d'espèces dont la chasse est autorisée, indiquez ici les références de l'autorisation de prélèvement accordée pour effectuer la capture des animaux dans le milieu naturel (article L. 424-11 du Code de l'environnement) :",
    "Justification": "Justification scientifique montrant que l'objectif de la procédure expérimentale ne peut être atteint en utilisant un animal élevé en vue d'une utilisation dans des procédures expérimentales :",
    "AnimauxGenetiquementAlteres": " - Animaux génétiquement altérés : ",
    "OuiAnimauxGenetiquementAlteres": "Si animaux génétiquement altérés :",
    "OuiAnimauxNonDomNonCaptifs": "Si animaux non domestiques non tenus en captivité :",
    "AnimauxGenetiquementModifiesEtSoucheMutante": "Animaux génétiquement modifiés ET Souche mutante autre",
    "AnimauxGenetiquementModifies": " - Animaux génétiquement modifiés",
    "AnimauxGenetiquementModifies2": " - Animaux génétiquement modifiés",
    "CreationLignee": "Création d'une lignée : ",
    "MaintienLigneeEtablie": "Maintien d'une lignée établie et/ou utilisation : ",
    "PhenotypeDommageable": "Phénotype dommageable : ",
    "NumeroRecipisse": "Numéro d’agrément, le cas échéant : ",
    "SoucheMutante2": " - Souche mutante autre",
    "CreationMutant": "Création d'un mutant : ",
    "MaintienLigneeEtablieMutant": "Maintien d'une lignée établie et/ou utilisation : ",
    "PhenotypeDommageableMutant": "Phénotype dommageable : ",
    "OrigineAnimaux": "3.4.9. Origine des animaux tenus en captivité :",
    "ElevesAFinExperimentale": "Les animaux destinés à être utilisés dans les procédures expérimentales appartenant aux espèces dont la liste est fixée réglementairement sont-ils élevés à cette fin et proviennent-ils d'éleveurs ou de fournisseurs agréés ?",
    "OuiEleves": "", # This is a placeholder, actual label depends on count
    "OuiElevesAFinExperimentale": "Nom de l'établissement :",
    "NomEtablissement": " Nom de l'établissement :",
    "AdressePostaleEtablissement": " Adresse postale",
    "AnimauxFournis": " Animaux Fournis : ",
    "NonEleves": "Si non, justifier scientifiquement l'utilisation d'animaux qui ne proviennent pas d'éleveurs ou de fournisseurs agréés :",
    "VotreEUFournitToutOuPartieAnimaux": "Votre propre établissement utilisateur fournit-il tout ou partie des animaux du projet ?",
    "AutreEUFournitToutOuPartieAnimaux": "Un autre établissement utilisateur fournit-il tout ou partie des animaux du projet ?",
    "EtablissementsNonAgreesFournissantAnimaux": "Nom de l'établissement :", # This is a placeholder, actual label depends on count
    "NomEtablissementsNonAgreesFournissantAnimaux": "Nom de l'établissement :",
    "EtablissementEtatMembre": "Informations sur l'établissement :",
    "NomEtablissementsEtatsMembresFournissantAnimaux": "Nom de l'établissement :",
    "PaysEtablissementsEtatsMembresFournissantAnimaux": "Pays de l'établissement :",
    "EtablissementEtatTiers": "Informations sur l'établissement :",
    "NomEtablissementsEtatsTiersFournissantAnimaux": "Nom de l'établissement :",
    "PaysEtablissementsEtatsTiersFournissantAnimaux": "Pays de l'établissement :",
    "AnimauxReutilisesProjetPrecedent": "Les animaux sont-ils des animaux réutilisés d'un projet précédent ?",
    "AnimauxUtilises": "", # This is a placeholder, actual label depends on count
    "NombreAnimauxUtilises": "3.4.10. Nombre estimé d'animaux utilisés dans le projet :",
    "JustificationUtilisationEspeces": "Justification de ce nombre pour chacune des espèces animales utilisées :",
    "UtilisationQuelStade": "3.4.11. Indiquez à quel(s) stade(s) de développement les animaux seront utilisés et le justifier :",
    "SexeAnimauxUtilisesJustification": "3.4.12. Indiquez le sexe des animaux utilisés et le justifier :",
    "ProceduresExperimentales": "4. Procédures expérimentales",
    "ObjetsVises": "4.1. Objets visés par les procédures expérimentales :",
    "PointA": " - Point A",
    "PointB": " - Point B",
    "PointC": " - Point C",
    "PointD": " - Point D",
    "PointE": " - Point E",
    "PointF": " - Point F",
    "PointG": " - Point G",
    "ExplicationsProcedures": "4.2 Nombre de procédures expérimentales :", # This is a placeholder, actual label depends on count
    "Procedure": "4.2.x Procédure x", # This will be dynamically numbered in the template
    "NomProcedure": "Nom de la procédure :",
    "ClassificationProcedure": "     Proposition de classification de la procédure selon le degré de sévérité : ",
    "DescriptionDetaillee": "     Description détaillée de la procédure expérimentale ",
    "PertinenceJustification": "     Pertinence et justification de la procédure expérimentale : ",
    "NombreLots": "      Indiquez le nombre de lots et le nombre d'animaux par lots, et les justifier :",
    "PointsLimitesAdaptes": "      Indiquez pour chaque espèce les points limites adaptés, suffisamment prédictifs et précoces pour permettre de limiter la douleur à son minimum, sans remettre en cause les résultats du projet : ",
    "PrelevementEtFrequence": "      Indiquez le cas échéant le prélèvement, ainsi que la fréquence et le(s) volume(s) prélevés : ",
    "MethodeSuppressionDouleur": "     Indiquez le cas échéant les méthodes pour réduire ou supprimer la douleur, la souffrance et l'angoisse [liste des médications . anesthésiques, analgésiques, anti-inflammatoires, en précisant les doses, voies, durées et fréquences d'administration], y compris le raffinement des conditions d'hébergement, d'élevage et de soins :",
    "MethodeSuppressionSouffrance": "      Indiquez le cas échéant les dispositions prises en vue de réduire, d'éviter et d'atténuer toute forme de souffrance des animaux de la naissance à la mort : ",
    "RaisonsScientifiquesDerogationAnestesie": "     Indiquez le cas échéant les raisons scientifiques justifiant une dérogation à l'anesthésie des animaux : ",
    "RaisonsScientifiquesDerogationHebergement": "     Indiquez le cas échéant les raisons scientifiques justifiant une dérogation aux conditions d'hébergement des animaux : ",
    "DispositionsPrisesEviterDoubleEmploi": "     Dispositions prises pour éviter tout double emploi injustifié des procédures expérimentales, le cas échéant : ",
    "DevenirAnimaux": "     Devenir des animaux à la fin de cette procédure expérimentale : ",
    "MiseAMortAnimaux": "      mise à mort ? : ",
    "AnimauxMisAMort": "	précisez les animaux concernés :",
    "GardeEnVie": "      animal gardé en vie ? : ",
    "AnimauxGardesEnVie": "	précisez les animaux concernés :",
    "DecisionVeterinaire": "Précisez si la décision a été prise par le vétérinaire ou toute autre personne compétente désignée par le responsable du projet:",
    "MiseEnLiberte": "     Placement ou mise en liberté des animaux ? : ",
    "AnimauxMisEnLiberte": "Précisez les animaux concernés:",
    "UtilisationAnimauxProjetAnterieur": "4.3. Si le projet utilise des animaux réutilisés d'un projet antérieur :",
    "GraviteReelle": "Gravité réelle de la (ou des) procédure(s) antérieure(s) :",
    "GraviteLegere": "- légère : ",
    "GraviteModeree": "- modérée : ",
    "GraviteSevere": "- sévère : ",
    "ElementsScientifiquesJustifiantDerogation": "Si des animaux sont issus d’une procédure antérieure « sévère », apporter les éléments scientifiques justifiant la dérogation, pour autant que ces animaux n’aient pas été utilisés plus d’une fois dans une procédure expérimentale entrainant une douleur intense, de l’angoisse ou une souffrance équivalente. Dans ce cas, le MESRI peut autoriser la réutilisation, après avis du comité d’éthique en expérimentation animale dont relève l’établissement ",
    "EffetCumulatif": "Effet cumulatif de cette réutilisation sur les animaux :",
    "AvisVeterinaireFavorable": "L'avis vétérinaire est-il favorable en prenant en considération le sort de l'animal concerné sur toute sa durée de vie ?",
    "EtatDeSanteRecouvre": "L'animal réutilisé a-t-il pleinement recouvré son état de santé et de bien-être général ?",
    "ProjetContenantSouffranceSevere": "4.4. Cas particulier des projets contenant une procédure expérimentale impliquant une douleur, une angoisse ou une souffrance sévère et susceptible de se prolonger sans qu'il soit possible de les soulager",
    "DerogationHebergement": "4.5. Dérogation aux conditions d'hébergement",
    "RaisonsScientifiques": "Raisons scientifiques :",
    "ProcedureEnCause": "De quelle(s) procédure(s) expérimentale(s) du projet s'agit-il ?",
    "JustificationScientifiqueProcedure": "Justifiez scientifiquement les raisons à l'origine d'une demande de dérogation :",
    "PublishNtsProjectRequest": "5. Résumé au format européen",
    "country": "Pays :",
    "language": "Langue :",
    "euSubmission": "Soumission européenne :",
    "ntsNationalId": "Identifiant national NTS :",
    "projectTitle": "5.1. Intitulé du projet [Repris automatiquement du champ 1.2]",
    "projectDuration": "5.2. Durée du projet (en mois) [Repris automatiquement du champ 1.3]",
    "keywords": "5.3. Mots-clés ",
    "keyword": "- Mot-clé n1 : ",
    "keyword2": "- Mot-clé n2 : ",
    "keyword3": "- Mot-clé n3 : ",
    "keyword4": "- Mot-clé n4 : ",
    "keyword5": "- Mot-clé n5 : ",
    "keyword6": "- Mot-clé n6 : ",
    "projectPurposes": "5.4. Finalités du projet. Sélectionner dans la liste proposée la ou les finalités du projet",
    "purpose": "- Finalité : ",
    "objectivesAndBenefits": "5.5. Objectifs et bénéfices escomptés du projet ",
    "projectObjectives": "5.5.1. Décrire les objectifs du projet [GRAND PUBLIC] ",
    "potentialBenefits": "5.5.2. Quels sont les bénéfices susceptibles de découler de ce projet [GRAND PUBLIC]?",
    "predictedHarms": "5.6. Nuisances prévues ",
    "procedures": "5.6.1. À quels types d’interventions les animaux seront-ils soumis (par exemple, prélèvements sur animaux vigiles, procédures chirurgicales) ? Indiquer leur nombre et leur durée. ",
    "expectedImpacts": "5.6.2. Quels sont les effets / ou effets indésirables prévus sur les animaux ? ",
    "expectedHarms": "5.7. Quelles espèces est-il prévu d’utiliser? Quels sont le degré de gravité des procédures et le nombre d’animaux prévus dans chaque catégorie de gravité (par espèce) ? ",
    "harm": "Gravité par espèce :",
    "species": "- Espèce : ",
    "nonRecovery": "     Sans réveil : ",
    "mild": "     Légère : ",
    "moderate": "     Modérée : ",
    "severe": "     Sévère : ",
    "fateOfAnimalsKeptAlive": "5.8. Qu’adviendra-t-il des animaux maintenus en vie à la fin du projet ?",
    "fate": "Sort des animaux par espèce :",
    "reused": "     Réutilisés : ",
    "returned": "     Replacés : ",
    "rehomed": "     Adoptés : ",
    "fateReasons": "5.9. Justifier le sort prévu de tous les animaux à l'issue de chaque procédure ",
    "applicationOfTheThreeRs": "5.10. Application de la règle des «trois R» ",
    "replacement": "1. Remplacement. [Repris automatiquement du champ 3.3.6.1] ",
    "reduction": "2. Réduction. [Repris automatiquement du champ 3.3.6.2] ",
    "refinement": "3. Raffinement. [Repris automatiquement du champ 3.3.6.3] ",
    "speciesChoiceExplanation": "5.11. Expliquer le choix des espèces et les stades de développement y afférents. [Repris automatiquement des champs 3.4.3 et 3.4.11, dans la limite des 2500 premiers caractères]",
}

def parse_xml_to_structured_data(xml_content):
    """
    Recursively parses XML content into a structured dictionary, preserving hierarchy
    and applying labels from FORM_LABEL_MAP.
    """
    parser = etree.XMLParser(recover=True, encoding='utf-8')
    root = etree.fromstring(xml_content.encode('utf-8'), parser=parser)
    
    def _parse_element(element):
        data = OrderedDict()

        data['tag'] = element.tag # Original tag for reconstruction
        # Use original tag for FORM_LABEL_MAP lookup (not attribute-modified)
        data['label'] = FORM_LABEL_MAP.get(element.tag, element.tag.replace('_', ' '))
        # Store attributes separately for display/reconstruction
        data['attributes'] = dict(element.attrib) if element.attrib else {}

        value = element.text.strip() if element.text and element.text.strip() else None
        if value is not None:
            if value.lower() == 'true':
                data['value'] = True
            elif value.lower() == 'false':
                data['value'] = False
            else:
                data['value'] = value
        else:
            data['value'] = None

        children_data = OrderedDict()
        for child in element:
            child_tag = child.tag
            if child_tag not in children_data:
                children_data[child_tag] = []
            children_data[child_tag].append(_parse_element(child))

        # Dynamic numbering for repeated elements
        if element.tag == 'ExplicationsProcedures' and 'Procedure' in children_data:
            for i, proc in enumerate(children_data['Procedure']):
                proc['label'] = f"4.2.{i+1} Procédure {i+1}"
        if element.tag == 'ResponsablesMiseEnOeuvre' and 'CoordonneesResponsablesMiseEnOeuvre' in children_data:
            for i, resp in enumerate(children_data['CoordonneesResponsablesMiseEnOeuvre']):
                resp['label'] = f"Responsable de la mise en oeuvre {i+1}"
        if element.tag == 'ResponsablesBienEtre' and 'CoordonneesResponsablesBienEtre' in children_data:
            for i, resp in enumerate(children_data['CoordonneesResponsablesBienEtre']):
                resp['label'] = f"Responsable du bien-être {i+1}"
        if element.tag == 'OuiEleves' and 'OuiElevesAFinExperimentale' in children_data:
            for i, etab in enumerate(children_data['OuiElevesAFinExperimentale']):
                etab['label'] = f"Établissement éleveur {i+1}"

        if children_data:
            data['children'] = children_data
        elif data['value'] is None:
            # Empty leaf elements: set value to empty string so they remain editable
            data['value'] = ''

        return data

    return _parse_element(root)

def merge_structured_data_for_form(data_list):
    """
    Merge multiple structured data instances into a single 'superset' structure.
    Used for FieldList form creation so the form class has fields for ALL possible
    sub-elements across all instances (e.g., different procedures may have different fields).
    """
    if not data_list:
        return None
    if len(data_list) == 1:
        return data_list[0]

    merged = OrderedDict()
    merged['tag'] = data_list[0]['tag']
    merged['label'] = data_list[0]['label']
    merged['attributes'] = data_list[0].get('attributes', {})
    # Use the first non-None/non-empty value
    merged['value'] = next((d['value'] for d in data_list if d.get('value')), data_list[0].get('value'))

    # Collect all children from all instances
    all_children_by_tag = OrderedDict()
    for data in data_list:
        if data.get('children'):
            for child_tag, children in data['children'].items():
                if child_tag not in all_children_by_tag:
                    all_children_by_tag[child_tag] = []
                all_children_by_tag[child_tag].extend(children)

    if all_children_by_tag:
        merged_children = OrderedDict()
        for child_tag, instances in all_children_by_tag.items():
            # Recursively merge all instances of this child tag into one representative
            merged_child = merge_structured_data_for_form(instances)
            merged_children[child_tag] = [merged_child]
        merged['children'] = merged_children

    return merged


def flatten_xml_to_form_data(xml_content):
    form_data = OrderedDict()
    try:
        parser = etree.XMLParser(recover=True, encoding='utf-8')
        root = etree.fromstring(xml_content.encode('utf-8'), parser=parser)
        
        for element in root.iter():
            if not list(element): # Only process leaf elements (no children)
                path = element.getroottree().getpath(element)
                
                # Replace lxml paths `/*[1]` by `TagName[1]` for consistency
                # This regex needs to be more robust for nested indexed elements
                # For now, keep it as is, but note it's a potential source of issues
                path = re.sub(r'/\*\[(\d+)\]', lambda m: f'/{root.tag}[{m.group(1)}]', path)
                path = re.sub(r'/\*\[(\d+)\]/(\w+)', r'/\2[\1]', path)
                
                value = ''
                if element.text and element.text.strip():
                    raw_value = element.text.strip()
                    if raw_value.lower() == 'true': value = True
                    elif raw_value.lower() == 'false': value = False
                    else: value = raw_value
                form_data[path] = value
    except etree.XMLSyntaxError as e:
        raise ValueError(f"Erreur de syntaxe XML: {e}")
    return form_data

def reconstruct_xml_from_form_data(original_xml_content, form_data):
    try:
        parser = etree.XMLParser(recover=True, encoding='utf-8')
        root = etree.fromstring(original_xml_content.encode('utf-8'), parser=parser)
        
        for path, value in form_data.items():
            try:
                # lxml's xpath method natively handles paths with indices
                elements = root.xpath(path)
                if elements:
                    elem_to_update = elements[0]
                    if isinstance(value, bool):
                        elem_to_update.text = 'true' if value else 'false'
                    else:
                        elem_to_update.text = str(value) if value is not None else ''
            except etree.XPathEvalError:
                print(f"Attention : Chemin XPath invalide ignoré lors de la reconstruction : {path}")

        return etree.tostring(root, encoding='utf-8', pretty_print=True).decode('utf-8')
    except etree.XMLSyntaxError as e:
        raise ValueError(f"Erreur de syntaxe XML lors de la reconstruction: {e}")

def populate_nested_form_from_structured_data(form, structured_data):
    """
    Recursively populates a WTForms form with data from a structured dictionary.
    """
    # If the current structured_data has a value and no children, it's a leaf node
    if structured_data.get('value') is not None and not structured_data.get('children'):
        field_name = structured_data['tag']
        if hasattr(form, field_name):
            field = getattr(form, field_name)
            field.data = structured_data['value']
    
    # If it has children, iterate through them
    if structured_data.get('children'):
        for child_tag, children_list in structured_data['children'].items():
            if hasattr(form, child_tag):
                form_field = getattr(form, child_tag)
                if isinstance(form_field, FieldList):
                    # For FieldList, populate each entry
                    for i, child_data in enumerate(children_list):
                        if i < len(form_field.entries):
                            populate_nested_form_from_structured_data(form_field.entries[i].form, child_data)
                elif isinstance(form_field, FormField):
                    # For FormField, populate its nested form
                    populate_nested_form_from_structured_data(form_field.form, children_list[0])

def reconstruct_xml_from_structured_form_data(original_structured_data, form):
    """
    Recursively reconstructs XML string from the submitted nested form data,
    preserving original structure and attributes from original_structured_data.
    """
    root_tag = original_structured_data['tag']
    root_element = etree.Element(root_tag)

    def _reconstruct_element(current_xml_element, current_structured_data, current_form_data):
        # When form_data is None, preserve original values from structured_data
        if current_form_data is None:
            if current_structured_data.get('value') is not None and not current_structured_data.get('children'):
                value = current_structured_data['value']
                if isinstance(value, bool):
                    current_xml_element.text = 'true' if value else 'false'
                elif value:
                    current_xml_element.text = str(value)
            if current_structured_data.get('children'):
                for child_tag, children_list in current_structured_data['children'].items():
                    for child_data in children_list:
                        new_child_element = etree.SubElement(current_xml_element, child_tag)
                        _reconstruct_element(new_child_element, child_data, None)
            return

        # Set value for leaf nodes
        if current_structured_data.get('value') is not None and not current_structured_data.get('children'):
            field_name = current_structured_data['tag']
            if hasattr(current_form_data, field_name):
                value = getattr(current_form_data, field_name).data
                if isinstance(value, bool):
                    current_xml_element.text = 'true' if value else 'false'
                else:
                    current_xml_element.text = str(value) if value is not None else ''

        # Recursively process children
        if current_structured_data.get('children'):
            for child_tag, children_list in current_structured_data['children'].items():
                if current_form_data is not None and hasattr(current_form_data, child_tag):
                    form_field = getattr(current_form_data, child_tag)
                    if isinstance(form_field, FieldList):
                        for i, child_form_entry in enumerate(form_field.entries):
                            if i < len(children_list):
                                new_child_element = etree.SubElement(current_xml_element, child_tag)
                                _reconstruct_element(new_child_element, children_list[i], child_form_entry.form)
                    elif isinstance(form_field, FormField):
                        new_child_element = etree.SubElement(current_xml_element, child_tag)
                        _reconstruct_element(new_child_element, children_list[0], form_field.form)
                else: # If a child element exists in structured_data but not in form
                    for child_data in children_list:
                        new_child_element = etree.SubElement(current_xml_element, child_tag)
                        _reconstruct_element(new_child_element, child_data, None)

    _reconstruct_element(root_element, original_structured_data, form)
    return etree.tostring(root_element, encoding='utf-8', pretty_print=True).decode('utf-8')


def _get_heading_depth(label):
    """Determine heading depth from numbered label prefix.
    '3.' → 1, '3.1.' → 2, '3.3.2.' → 3, '3.3.6.1.' → 4
    """
    if not label:
        return 0
    match = re.match(r'^(\d+\.)+', label)
    if match:
        return match.group().count('.')
    return 0


def _add_heading(doc, text, level):
    doc.add_heading(text, level=min(level, 4))


def _add_paragraph(doc, title, text, bold_title=True):
    if not text:
        return
    p = doc.add_paragraph()
    if title:
        run = p.add_run(f"{title} ")
        if bold_title:
            run.bold = True
    text_to_add = str(text)
    if isinstance(text, str):
        if text.lower() == 'true':
            text_to_add = 'Oui'
        elif text.lower() == 'false':
            text_to_add = 'Non'
    p.add_run(text_to_add)


def _add_multiline_paragraph(doc, title, text, bold_title=True):
    """Add a paragraph with title and multiline text, preserving line breaks."""
    if not text:
        return
    lines = text.split('\n')
    p = doc.add_paragraph()
    if title:
        run = p.add_run(f"{title} ")
        if bold_title:
            run.bold = True
    # Add first line to the same paragraph
    first_line = lines[0].strip() if lines else ''
    if first_line:
        p.add_run(first_line)
    # Remaining lines as new paragraphs
    for line in lines[1:]:
        clean_line = line.strip()
        if clean_line:
            doc.add_paragraph(clean_line)


def generate_docx_from_xml(xml_content):
    try:
        structured_data = parse_xml_to_structured_data(xml_content)
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        def _add_content_to_docx(data, level=0):
            label = data.get('label')
            value = data.get('value')
            children = data.get('children')
            tag = data.get('tag', '')
            is_leaf = value is not None and not children

            depth = _get_heading_depth(label)
            heading_rendered = False

            if label and label != tag:
                if depth > 0:
                    # Numbered section: heading level based on depth
                    heading_level = min(depth, 4)
                    if not is_leaf:
                        # Section header only (not a leaf field)
                        _add_heading(doc, label, level=heading_level)
                        heading_rendered = True
                elif label.startswith(' - ') or label.startswith('     '):
                    # Indented items - render inline with value
                    pass
                elif not is_leaf and children:
                    # Non-numbered parent with children: sub-heading
                    _add_heading(doc, label, level=min(level + 2, 4))
                    heading_rendered = True

            if is_leaf:
                if isinstance(value, bool):
                    display_label = label if label != tag else tag
                    _add_paragraph(doc, display_label, 'Oui' if value else 'Non')
                elif value and isinstance(value, str) and '\n' in value:
                    display_label = label if label != tag else tag
                    if depth > 0 and not heading_rendered:
                        _add_heading(doc, label, level=min(depth, 4))
                    _add_multiline_paragraph(doc, display_label if not heading_rendered else None, value)
                elif value:
                    display_label = label if label != tag else tag
                    if depth > 0 and not heading_rendered:
                        _add_heading(doc, label, level=min(depth, 4))
                        _add_paragraph(doc, None, value, bold_title=False)
                    else:
                        _add_paragraph(doc, display_label, value)

            if children:
                for child_tag, children_list in children.items():
                    for child in children_list:
                        _add_content_to_docx(child, level + 1)

            # Add a page break after top-level numbered sections
            if depth == 1 and level <= 1:
                doc.add_page_break()

        _add_content_to_docx(structured_data)

        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io
    except Exception as e:
        raise IOError(f"Erreur lors de la génération du DOCX: {e}")
